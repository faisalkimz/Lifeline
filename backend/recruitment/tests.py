
import os
from django.test import TestCase
from recruitment.resume_parser import ResumeParser
import docx

class ResumeParserTests(TestCase):
    def setUp(self):
        self.test_file = "test_resume_django.docx"
        # Create a dummy docx
        doc = docx.Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("john.doe@test.com | 0700123456")
        doc.add_paragraph("Skills: Python, Django")
        doc.save(self.test_file)

    def tearDown(self):
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

    def test_extract_name_email(self):
        parser = ResumeParser(self.test_file)
        data = parser.parse()
        self.assertIn("John Doe", data['name'])
        self.assertEqual(data['email'], "john.doe@test.com")
        self.assertIn("Python", data['skills'])
